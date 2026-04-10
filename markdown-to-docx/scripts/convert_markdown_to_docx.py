#!/usr/bin/env python3
"""
Convert a Markdown file with Mermaid diagrams to DOCX.

Usage:
  python scripts/convert_markdown_to_docx.py <input.md> [output.docx]

  input.md    - Path to the Markdown file to convert (required)
  output.docx - Path for the generated DOCX (optional; defaults to <input>.docx)

Steps:
  1. Pre-process markdown:
       - Strip manually-written '## Table of Contents' section (if present)
       - Inject TOC placeholders right after the first '---' separator
         (i.e. after the document header block: title + metadata + ---)
  2. Find ```mermaid blocks, render each to high-resolution PNG via mmdc
  3. Replace each mermaid block with a ![](image.png) reference
  4. Run pandoc (no --toc flag — TOC is inserted via python-docx instead)
  5. Post-process docx:
       - Replace TOC placeholders with a native Word TOC field
         ({ TOC \\o "1-3" \\h \\z \\u }) positioned after the header block
       - Apply table borders + header row formatting
"""

import re
import subprocess
import sys
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if len(sys.argv) < 2:
    print("Usage: python scripts/convert_markdown_to_docx.py <input.md> [output.docx]", file=sys.stderr)
    sys.exit(1)

INPUT_MD = Path(sys.argv[1])
if not INPUT_MD.exists():
    print(f"Error: input file '{INPUT_MD}' not found.", file=sys.stderr)
    sys.exit(1)

OUTPUT_DOCX = Path(sys.argv[2]) if len(sys.argv) >= 3 else INPUT_MD.with_suffix(".docx")
IMG_DIR = INPUT_MD.parent / "_mermaid_imgs"
MMDC = "mmdc"
CHROMIUM = "/usr/bin/chromium-browser"
PUPPETEER_CFG = Path.home() / "puppeteer.json"

# Mermaid render settings — high resolution for clear print/zoom quality
MMDC_WIDTH = 2400   # viewport width in pixels
MMDC_SCALE = 3      # output scale multiplier (effective: 7200px wide at 3×)

# Table style colours
HEADER_BG = "1F3864"
HEADER_FG = "FFFFFF"
BORDER_COLOR = "BFBFBF"
ALT_ROW_BG = "EBF0FA"

# Unique placeholder strings injected into the markdown then replaced in the DOCX.
# These must not appear anywhere in normal document content.
_TOC_HEADING_PLACEHOLDER = "WORD_TOC_HEADING_XZQY_PLACEHOLDER"
_TOC_FIELD_PLACEHOLDER   = "WORD_TOC_FIELD_XZQY_PLACEHOLDER"


# ---------------------------------------------------------------------------
# Markdown pre-processing
# ---------------------------------------------------------------------------

def preprocess_markdown(md_text: str) -> str:
    """Prepare markdown for pandoc conversion.

    1. Remove any manually-written '## Table of Contents' section.
    2. Inject TOC placeholders after the first '---' separator so the Word
       TOC field lands after the document header block (title + metadata),
       not at the very top of the document.
    """
    # Step 1: strip manual ToC (safe to apply even when absent)
    cleaned = re.sub(
        r"## Table of Contents\n.*?(?=\n## )",
        "",
        md_text,
        flags=re.DOTALL,
    )

    # Step 2: inject placeholders after the FIRST standalone '---' line.
    # That separator marks the end of the document header block.
    cleaned = re.sub(
        r"^---$",
        f"---\n\n{_TOC_HEADING_PLACEHOLDER}\n\n{_TOC_FIELD_PLACEHOLDER}",
        cleaned,
        count=1,
        flags=re.MULTILINE,
    )
    return cleaned


# ---------------------------------------------------------------------------
# Mermaid rendering
# ---------------------------------------------------------------------------

def render_mermaid_blocks(md_text: str, img_dir: Path) -> str:
    img_dir.mkdir(exist_ok=True)
    pattern = re.compile(r"```mermaid\n(.*?)```", re.DOTALL)
    counter = 0

    def replace_block(match):
        nonlocal counter
        counter += 1
        diagram_src = match.group(1).strip()
        mmd_file = img_dir / f"diagram_{counter:02d}.mmd"
        png_file = img_dir / f"diagram_{counter:02d}.png"
        mmd_file.write_text(diagram_src, encoding="utf-8")

        cmd = [
            MMDC,
            "-i", str(mmd_file),
            "-o", str(png_file),
            "-b", "white",
            "-w", str(MMDC_WIDTH),
            "--scale", str(MMDC_SCALE),
            "--puppeteerConfigFile", str(PUPPETEER_CFG),
        ]
        env = os.environ.copy()
        env["PUPPETEER_EXECUTABLE_PATH"] = CHROMIUM

        result = subprocess.run(cmd, capture_output=True, text=True, env=env)
        if result.returncode != 0:
            print(f"  [WARN] diagram_{counter:02d} failed: {result.stderr[:200]}", file=sys.stderr)
            return match.group(0)

        print(f"  [OK]   diagram_{counter:02d}.png  ({png_file.stat().st_size // 1024} KB)")
        return f"![]({png_file})\n"

    return pattern.sub(replace_block, md_text)


# ---------------------------------------------------------------------------
# TOC insertion (python-docx post-processing)
# ---------------------------------------------------------------------------

def _fill_toc_field(p) -> None:
    """Replace all content in paragraph element *p* with a Word TOC field.

    The field instruction ' TOC \\o "1-3" \\h \\z \\u ' produces a clickable,
    auto-updating Table of Contents covering heading levels 1–3.
    Word will show 'Error! No table of contents entries found.' until the user
    clicks 'Update Table' — this is expected behaviour for a dirty TOC field.
    """
    for child in list(p):
        p.remove(child)

    def run_with(child_el) -> None:
        r = OxmlElement("w:r")
        r.append(child_el)
        p.append(r)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")   # tells Word to update on open
    run_with(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_with(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_with(fld_end)


def insert_toc_after_header(doc: Document) -> None:
    """Find the two TOC placeholder paragraphs and replace them in-place.

    _TOC_HEADING_PLACEHOLDER → bold 'Table of Contents' paragraph
    _TOC_FIELD_PLACEHOLDER   → native Word { TOC \\o "1-3" \\h \\z \\u } field
    """
    from docx.shared import Pt

    found = 0
    for para in doc.paragraphs:
        text = para.text.strip()

        if text == _TOC_HEADING_PLACEHOLDER:
            p = para._p
            for child in list(p):
                p.remove(child)
            pPr = OxmlElement("w:pPr")
            p.append(pPr)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            bold = OxmlElement("w:b")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "28")        # 14 pt  (half-points)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), "28")
            rPr.append(bold)
            rPr.append(sz)
            rPr.append(szCs)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = "Table of Contents"
            r.append(t)
            p.append(r)
            found += 1

        elif text == _TOC_FIELD_PLACEHOLDER:
            _fill_toc_field(para._p)
            found += 1

    if found == 2:
        print("  [OK]   Native Word TOC field inserted after document header")
    else:
        print(f"  [WARN] Expected 2 TOC placeholders, found {found}. TOC may be missing.", file=sys.stderr)


# ---------------------------------------------------------------------------
# Table formatting (python-docx post-processing)
# ---------------------------------------------------------------------------

def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top=BORDER_COLOR,
                    bottom=BORDER_COLOR,
                    left=BORDER_COLOR,
                    right=BORDER_COLOR,
                )
                if is_header:
                    set_cell_bg(cell, HEADER_BG)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif row_idx % 2 == 0:
                    set_cell_bg(cell, ALT_ROW_BG)

    print(f"  [OK]   Table formatting applied ({len(doc.tables)} tables)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    input_md = INPUT_MD.resolve()
    output_docx = OUTPUT_DOCX.resolve()
    img_dir = input_md.parent / "_mermaid_imgs"

    os.chdir(input_md.parent)
    PUPPETEER_CFG.write_text(
        '{"args":["--no-sandbox","--disable-setuid-sandbox"]}',
        encoding="utf-8",
    )

    print(f"Reading {input_md} ...")
    md_text = input_md.read_text(encoding="utf-8")

    print("Pre-processing markdown (strip manual ToC, inject TOC placeholders) ...")
    md_text = preprocess_markdown(md_text)

    print("Rendering Mermaid diagrams ...")
    modified_md = render_mermaid_blocks(md_text, img_dir)

    tmp_md = input_md.parent / f"_{input_md.stem}_rendered.md"
    tmp_md.write_text(modified_md, encoding="utf-8")

    print(f"Running pandoc → {output_docx} ...")
    cmd = [
        "pandoc", str(tmp_md),
        "-o", str(output_docx),
        "--from", "markdown",
        "--to", "docx",
        # No --toc flag: TOC is inserted by python-docx after the header block,
        # not by pandoc which would place it at the very beginning of the document.
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"[ERROR] pandoc failed:\n{result.stderr}", file=sys.stderr)
        sys.exit(1)

    tmp_md.unlink()

    print("Inserting native Word TOC field after document header ...")
    doc = Document(str(output_docx))
    insert_toc_after_header(doc)

    print("Applying table formatting ...")
    format_tables(doc)
    doc.save(str(output_docx))

    print(f"\nDone → {output_docx} ({output_docx.stat().st_size // 1024} KB)")
    print()
    print("Note: Open the .docx in Word and click 'Update Table' when prompted")
    print("      to populate the Table of Contents with page numbers.")


if __name__ == "__main__":
    main()
